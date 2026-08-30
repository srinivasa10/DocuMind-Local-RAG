import csv
import json
import pytest
from pathlib import Path

from app.rag.loaders import extract_text, _extract_csv_tsv, _extract_json, _extract_plain


def test_plain_text_and_markdown(tmp_path: Path):
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("Hello World\nLine 2", encoding="utf-8")
    res = extract_text(txt_file)
    assert len(res) == 1
    assert "Hello World" in res[0].text

    md_file = tmp_path / "readme.md"
    md_file.write_text("# Heading\nSome markdown content", encoding="utf-8")
    res_md = extract_text(md_file)
    assert len(res_md) == 1
    assert "Heading" in res_md[0].text


def test_csv_extraction(tmp_path: Path):
    csv_file = tmp_path / "data.csv"
    with open(csv_file, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Role", "Department"])
        writer.writerow(["Srinivasa", "Programmer Analyst", "AI/ML"])
    
    res = extract_text(csv_file)
    assert len(res) == 1
    assert "Name | Role | Department" in res[0].text
    assert "Srinivasa | Programmer Analyst | AI/ML" in res[0].text


def test_json_extraction(tmp_path: Path):
    json_file = tmp_path / "data.json"
    data = {"candidate": "Srinivasa", "skills": ["Python", "React", "FastAPI"]}
    json_file.write_text(json.dumps(data), encoding="utf-8")

    res = extract_text(json_file)
    assert len(res) == 1
    assert "Srinivasa" in res[0].text
    assert "FastAPI" in res[0].text


def test_docx_extraction(tmp_path: Path):
    import docx

    doc_file = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading("Resume Summary", level=1)
    doc.add_paragraph("Experienced full-stack engineer.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"
    doc.save(str(doc_file))

    res = extract_text(doc_file)
    assert len(res) == 1
    assert "Resume Summary" in res[0].text
    assert "Experienced full-stack engineer." in res[0].text
    assert "Skill | Level" in res[0].text
    assert "Python | Expert" in res[0].text


def test_excel_extraction(tmp_path: Path):
    import openpyxl

    wb_file = tmp_path / "test.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employees"
    ws.append(["ID", "Name", "Score"])
    ws.append([101, "Srinivasa", 99.5])
    wb.save(str(wb_file))

    res = extract_text(wb_file)
    assert len(res) == 1
    assert "Sheet: Employees" in res[0].text
    assert "ID | Name | Score" in res[0].text
    assert "101 | Srinivasa | 99.5" in res[0].text
